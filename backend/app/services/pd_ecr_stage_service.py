"""
Stage an uploaded file for human review before final ingestion.

Flow:
1. Save original file → uploads/
2. Generate preview PDF (Excel/Word → PDF; PDF stays as-is)
3. Parse text + extract metadata, sections, tables
4. Store as PdEcrStagedDocument (status=pending)
5. User reviews & edits → confirms → creates PdEcrCase + VecStoredChunk
"""

from __future__ import annotations

import json
import logging
import re
import shutil
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from sqlmodel import Session, select

from app.models import (
    HistoricalSourceDocument,
    PdEcrCase,
    PdEcrModule,
    PdEcrStagedDocument,
    PdEcrStagedDocumentUpdate,
    User,
    PD_ECR_DEFAULT_MODULES,
)
from app.services.pd_ecr_audit_service import write_activity
from app.services.pd_ecr_import_service import (
    extract_metadata as _extract_db_metadata,
    safe_read_text,
    sanitize_case_no,
    KNOWLEDGE_DIR,
)

logger = logging.getLogger(__name__)

UPLOAD_DIR = Path(__file__).resolve().parents[1] / "uploads"


def _now_utc() -> datetime:
    return datetime.now(timezone.utc)


# ──────────────────────────────────────────────────────────────────────
# Section & table extraction
# ──────────────────────────────────────────────────────────────────────

_RE_HEADING = re.compile(r"^(#{1,6})\s+(.+)$")
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
_RE_TABLE_SEP = re.compile(r"^\|[\s\-:]+\|$")


def _parse_sections_and_tables(
    md_text: str,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Split markdown text into heading-delimited sections and extract tables.

    Returns (sections, tables) where each section is:
        {index, heading, level, content, page_no (inferred)}
    and each table is:
        {index, caption, headers: [...], rows: [[...]], page_no}
    """
    lines = md_text.splitlines()
    sections: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    current_heading = ""
    current_level = 0
    current_lines: list[str] = []
    section_idx = 0
    page_estimate = 1

    # Crude page-number inference: "## Page N" or blank line after ~40 lines
    _RE_PAGE_BREAK = re.compile(r"^##\s*Page\s+(\d+)", re.I)
    lines_since_page = 0

    def _flush_section():
        nonlocal section_idx
        content = "\n".join(current_lines).strip()
        if content or current_heading:
            sections.append({
                "index": section_idx,
                "heading": current_heading,
                "level": current_level,
                "content": content,
                "page_no": page_estimate,
            })
            section_idx += 1
        current_lines.clear()

    # Detect markdown tables
    table_buffer: list[str] = []
    in_table = False
    table_headers: list[str] = []

    def _flush_table():
        nonlocal in_table
        if table_headers and len(table_buffer) >= 1:
            tables.append({
                "index": len(tables),
                "caption": current_heading or "",
                "headers": table_headers,
                "rows": [
                    [c.strip() for c in row.strip("|").split("|")]
                    for row in table_buffer
                ],
                "page_no": page_estimate,
            })
        table_buffer.clear()
        table_headers.clear()
        in_table = False

    for line in lines:
        stripped = line.strip()

        # Page break detection
        pm = _RE_PAGE_BREAK.match(stripped)
        if pm:
            page_estimate = int(pm.group(1))
            lines_since_page = 0
            continue

        lines_since_page += 1
        if lines_since_page > 50:
            page_estimate += 1
            lines_since_page = 0

        # Heading
        hm = _RE_HEADING.match(stripped)
        if hm and _RE_TABLE_ROW.match(stripped) is None:
            if in_table:
                _flush_table()
            _flush_section()
            current_level = len(hm.group(1))
            current_heading = hm.group(2).strip()
            continue

        # Table row
        if _RE_TABLE_ROW.match(stripped):
            if not in_table:
                # Previous line might be table caption — if short and not empty
                if current_lines and len(current_lines[-1]) < 100:
                    pass  # keep as potential caption
                in_table = True
            if _RE_TABLE_SEP.match(stripped):
                if not table_headers and table_buffer:
                    # The row before separator was the header
                    table_headers = [
                        c.strip() for c in table_buffer[-1].strip("|").split("|")
                    ]
                    table_buffer.pop()
                continue
            table_buffer.append(stripped)
            continue
        else:
            if in_table:
                _flush_table()

        current_lines.append(line)

    # Flush any remaining content
    if in_table:
        _flush_table()
    _flush_section()

    return sections, tables


# ──────────────────────────────────────────────────────────────────────
# Parsing pipeline
# ──────────────────────────────────────────────────────────────────────

def _parse_file(file_path: Path, suffix: str) -> dict[str, Any]:
    """Parse a file using the structured PD-ECR extraction pipeline.

    Returns:
        parsed_text, parsed_by, metadata (flat dict for form display),
        structured (full template JSON), sections, tables
    """
    from app.rag.excel_to_markdown import convert_excel as _convert_excel
    from app.rag.pdf_to_markdown import convert_pdf as _convert_pdf
    from app.rag.pdecr_structured_extractor import extract_structured
    from app.rag.text_cleaner import clean_text

    KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)

    parsed_text = ""
    parsed_by = ""
    controls_json: list[dict[str, Any]] = []

    if suffix in (".xlsx", ".xlsm", ".xls"):
        if suffix in (".xlsx", ".xlsm"):
            try:
                from app.rag.xlsx_controls import extract_xlsx_controls

                controls_json = extract_xlsx_controls(file_path)
            except Exception:
                logger.warning("XLSX control extraction failed", exc_info=True)

        _convert_excel(file_path)
        md_path = KNOWLEDGE_DIR / f"{file_path.stem}.md"
        parsed_by = "excel_to_markdown"

        # Excel→PDF→MinerU for richer structured extraction
        try:
            if shutil.which("mineru") is not None:
                from app.rag.excel_to_pdf import convert_excel_to_pdf
                from app.rag.pdf_to_markdown import run_mineru
                pdf_path = convert_excel_to_pdf(file_path)
                if pdf_path:
                    mineru_md = run_mineru(pdf_path, backend="pipeline")
                    mineru_text = mineru_md.read_text(encoding="utf-8", errors="ignore")
                    mineru_out = KNOWLEDGE_DIR / f"{file_path.stem}_mineru.md"
                    mineru_out.write_text(mineru_text, encoding="utf-8")
                    if mineru_text.strip():
                        parsed_text = mineru_text
                        parsed_by = "excel_to_pdf_to_mineru"
        except Exception:
            logger.warning("Excel→PDF→MinerU failed, using keyword-filtered", exc_info=True)

        if not parsed_text and md_path.exists():
            parsed_text = safe_read_text(md_path)

    elif suffix == ".pdf":
        _convert_pdf(file_path)
        md_path = KNOWLEDGE_DIR / f"{file_path.stem}.md"
        parsed_by = "pdf_to_markdown"
        if md_path.exists():
            parsed_text = safe_read_text(md_path)

    elif suffix in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(str(file_path))
            parsed_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            parsed_by = "python_docx"
        except ImportError:
            parsed_text = ""
        if not parsed_text:
            parsed_text = f"[Word document: {file_path.name}]"
            parsed_by = "fallback"

    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    cleaned = clean_text(parsed_text) if parsed_text else ""

    # ── Run template-aware structured extraction (LLM first, rules fallback) ──
    from app.rag.pdecr_llm_extractor import extract_with_llm_fallback

    structured = {}
    try:
        if cleaned:
            structured = extract_with_llm_fallback(cleaned, rule_based_extractor=extract_structured)
    except Exception:
        logger.warning("Structured extraction failed", exc_info=True)

    # ── Build flat metadata for the review form ──
    ident = structured.get("identification", {})
    change = structured.get("change_request", {})
    metadata = {
        "case_no": ident.get("dc_no", ""),
        "dc_no": ident.get("dc_no", ""),
        "mcr_no": ident.get("mcr_no", ""),
        "date": ident.get("date", ""),
        "customer_project": ident.get("customer_project", ""),
        "product_no": ident.get("product_no", ""),
        "part_no": ident.get("part_no", ""),
        "change_type": ident.get("change_type", ""),
        "sample_type": ident.get("sample_type", ""),
        "initiator": ident.get("initiator", ""),
        "reason": change.get("reason", ""),
        "change_source": change.get("change_source", ""),
        "change_proposal": change.get("change_proposal", ""),
        "title": ident.get("dc_no", file_path.stem),
        "controls_json": controls_json,
    }
    if controls_json:
        structured["_controls"] = controls_json

    # ── Extract sections & tables for display ──
    sections, tables = _parse_sections_and_tables(cleaned)

    return {
        "parsed_text": cleaned,
        "parsed_by": parsed_by,
        "metadata": metadata,
        "structured": structured,
        "sections": sections,
        "tables": tables,
    }


def _generate_preview_pdf(file_path: Path, suffix: str) -> Path | None:
    """Generate a PDF preview for non-PDF files. PDFs are copied as-is."""
    preview_dir = UPLOAD_DIR / "previews"
    preview_dir.mkdir(parents=True, exist_ok=True)

    preview_path = preview_dir / f"{file_path.stem}.pdf"

    if suffix == ".pdf":
        # PDF → copy as preview
        shutil.copy2(file_path, preview_path)
        return preview_path

    if suffix in (".xlsx", ".xlsm", ".xls"):
        from app.rag.excel_to_pdf import convert_excel_to_pdf
        pdf = convert_excel_to_pdf(file_path)
        if pdf:
            target = preview_dir / pdf.name
            shutil.move(str(pdf), str(target))
            return target
        return None

    if suffix in (".docx", ".doc"):
        try:
            import subprocess
            lo = shutil.which("libreoffice") or shutil.which("soffice")
            if lo:
                subprocess.run(
                    [lo, "--headless", "--convert-to", "pdf",
                     "--outdir", str(preview_dir), str(file_path)],
                    capture_output=True, timeout=60,
                )
                out = preview_dir / f"{file_path.stem}.pdf"
                if out.exists():
                    return out
        except Exception:
            logger.warning("Word→PDF conversion failed for %s", file_path.name, exc_info=True)
        return None

    return None


# ──────────────────────────────────────────────────────────────────────
# Public API used by the route handler
# ──────────────────────────────────────────────────────────────────────

def stage_uploaded_file(
    *,
    session: Session,
    file_path: Path,
    original_filename: str,
    user: User,
) -> PdEcrStagedDocument:
    """Parse an uploaded file and create a staged document for review."""
    suffix = file_path.suffix.lower()

    # ── Generate preview PDF ──
    preview_pdf_path = _generate_preview_pdf(file_path, suffix)

    # ── Parse ──
    parsed = _parse_file(file_path, suffix)

    # ── Create staged document ──
    staged = PdEcrStagedDocument(
        status="pending",
        original_filename=original_filename,
        original_file_path=str(file_path),
        preview_pdf_path=str(preview_pdf_path) if preview_pdf_path else None,
        file_type=suffix.lstrip("."),
        parsed_text=parsed["parsed_text"],
        metadata_json=parsed["metadata"],
        sections_json=parsed["sections"],
        tables_json=parsed["tables"],
        created_by_id=user.id,
    )
    session.add(staged)
    session.commit()
    session.refresh(staged)
    return staged


def get_staged_document(*, session: Session, doc_id: str) -> PdEcrStagedDocument | None:
    """Retrieve a staged document by ID."""
    try:
        uid = uuid.UUID(doc_id)
    except ValueError:
        return None
    return session.get(PdEcrStagedDocument, uid)


def update_staged_document(
    *,
    session: Session,
    doc: PdEcrStagedDocument,
    payload: PdEcrStagedDocumentUpdate,
) -> PdEcrStagedDocument:
    """Apply user edits to a staged document."""
    if payload.metadata_json is not None:
        doc.metadata_json = payload.metadata_json
    if payload.sections_json is not None:
        doc.sections_json = payload.sections_json
    if payload.tables_json is not None:
        doc.tables_json = payload.tables_json
    doc.updated_at = _now_utc()
    session.add(doc)
    session.commit()
    session.refresh(doc)
    return doc


def confirm_staged_document(
    *,
    session: Session,
    doc: PdEcrStagedDocument,
    user: User,
) -> dict[str, Any]:
    """Confirm a staged document: create PdEcrCase + modules + vector chunks.

    Returns the created case info and indexing status.
    """
    if doc.status == "confirmed":
        return {"status": "already_confirmed", "case_id": str(doc.confirmed_case_id)}

    metadata = doc.metadata_json
    case_no = sanitize_case_no(str(metadata.get("case_no") or f"PD-ECR-{doc.id.hex[:8]}"))

    # ── 1. Create or find PdEcrCase ──
    case = session.exec(
        select(PdEcrCase).where(PdEcrCase.case_no == case_no)
    ).first()

    is_new = case is None
    if is_new:
        case = PdEcrCase(
            case_no=case_no,
            title=str(metadata.get("title") or case_no)[:500],
            status="draft",
            source_type="file_upload",
            is_historical=False,
            created_by_id=user.id,
            owner_id=user.id,
            dc_no=metadata.get("dc_no"),
            mcr_no=metadata.get("mcr_no"),
            customer_project=metadata.get("customer_project"),
            product_no=metadata.get("product_no"),
            part_no=metadata.get("part_no"),
            change_type=metadata.get("change_type"),
            sample_type=metadata.get("sample_type"),
            initiator=metadata.get("initiator"),
        )
        session.add(case)
        session.flush()

        # Create default modules
        existing = set(
            session.exec(
                select(PdEcrModule.module_id).where(PdEcrModule.case_id == case.id)
            ).all()
        )
        for module_id, title in PD_ECR_DEFAULT_MODULES:
            if module_id in existing:
                continue
            session.add(PdEcrModule(
                case_id=case.id,
                module_id=module_id,
                title=title,
                content_json={},
                content_md="",
                source_cases=[case.case_no],
                source_files=[doc.original_filename],
                status="draft",
                updated_by_id=user.id,
            ))

    # Fill change-description module with the confirmed parsed text
    cd_module = session.exec(
        select(PdEcrModule).where(
            PdEcrModule.case_id == case.id,
            PdEcrModule.module_id == "change-description",
        )
    ).first()
    if cd_module:
        cd_module.content_md = doc.parsed_text[:20000]
        cd_module.content_json = {
            **(cd_module.content_json or {}),
            "source_preview": doc.parsed_text[:1000],
        }
        cd_module.source_files = sorted(set([*cd_module.source_files, doc.original_filename]))
        session.add(cd_module)

    # Fill basic-information module with metadata
    bi_module = session.exec(
        select(PdEcrModule).where(
            PdEcrModule.case_id == case.id,
            PdEcrModule.module_id == "basic-information",
        )
    ).first()
    if bi_module:
        bi_module.content_json = {k: v for k, v in metadata.items() if v}
        bi_module.source_files = sorted(set([*bi_module.source_files, doc.original_filename]))
        session.add(bi_module)

    # ── 2. Record source document ──
    session.add(HistoricalSourceDocument(
        case_id=case.id,
        imported_by_id=user.id,
        source_file=doc.original_filename,
        source_path=doc.original_file_path,
        source_kind=doc.file_type,
        content_hash="",  # computed separately
        extracted_metadata=_build_source_extracted_metadata(metadata, doc),
        import_warnings=[],
    ))

    # ── 3. Mark staged doc as confirmed ──
    doc.status = "confirmed"
    doc.confirmed_at = _now_utc()
    doc.confirmed_case_id = case.id
    doc.updated_at = _now_utc()
    session.add(doc)

    write_activity(
        session=session,
        action="case.file_confirmed",
        case_id=case.id,
        actor_id=user.id,
        target_id=str(case.id),
        metadata={
            "original_filename": doc.original_filename,
            "is_new_case": is_new,
            "sections_count": len(doc.sections_json),
            "tables_count": len(doc.tables_json),
        },
    )
    session.commit()

    # ── 4. Build vector chunks and trigger FAISS rebuild ──
    chunk_count = _build_vector_chunks(doc, case)

    # Trigger FAISS rebuild in background
    import threading
    import traceback as _tb
    def _rebuild():
        try:
            from app.rag.ingest import rebuild_index
            rebuild_index()
        except Exception:
            _tb.print_exc()

    threading.Thread(target=_rebuild, daemon=True).start()

    return {
        "status": "confirmed",
        "case_id": str(case.id),
        "case_no": case_no,
        "is_new_case": is_new,
        "chunks_created": chunk_count,
        "indexing": {
            "pending": True,
            "message": "知识库正在后台索引中，新数据将在数秒后可供检索。",
        },
    }


# ──────────────────────────────────────────────────────────────────────
# Vector chunk builder
# ──────────────────────────────────────────────────────────────────────

def _build_source_extracted_metadata(
    metadata: dict[str, Any],
    doc: PdEcrStagedDocument,
) -> dict[str, Any]:
    """Persist staged structured JSON with the source document record."""
    display_pdf_path = doc.preview_pdf_path or (
        doc.original_file_path if doc.file_type == "pdf" else None
    )
    return {
        **metadata,
        "original_filename": doc.original_filename,
        "original_file_path": doc.original_file_path,
        "display_pdf_path": display_pdf_path,
        "preview_pdf_path": display_pdf_path,
        "pdf_file": Path(display_pdf_path).name if display_pdf_path else "",
        "sections_json": doc.sections_json,
        "tables_json": doc.tables_json,
        "controls_json": metadata.get("controls_json", []),
    }


def _build_vector_chunks(doc: PdEcrStagedDocument, case: PdEcrCase) -> int:
    """Build row-level chunks from structured PD-ECR JSON with metadata-aware
    context. Each chunk represents one logical row (a field, an impact item,
    a document check, an approval role, etc.) and carries its section & field
    provenance — enabling precise metadata-aware RAG retrieval.
    """
    import pickle
    from app.rag.ingest.build_index import VECTOR_DIR as VS_DIR
    from app.rag.pdecr_structured_extractor import extract_structured, build_row_chunks

    VS_DIR.mkdir(parents=True, exist_ok=True)
    chunks_path = VS_DIR / f"chunks_{doc.id.hex}.pkl"

    # Re-run structured extraction on the confirmed parsed text to get
    # the latest structured JSON (user may have edited sections/tables)
    parsed = doc.parsed_text
    structured = {}
    try:
        structured = extract_structured(parsed) if parsed else {}
    except Exception:
        logger.warning("Structured extraction failed during chunk build", exc_info=True)

    # Build row-level chunks with metadata
    all_chunks = build_row_chunks(
        structured,
        source_file=doc.original_filename,
        file_id=str(doc.id),
    )

    metadata = getattr(doc, "metadata_json", {}) or {}
    base_metadata = {
        key: metadata.get(key)
        for key in (
            "dc_no",
            "mcr_no",
            "customer_project",
            "product_no",
            "part_no",
            "change_type",
            "sample_type",
            "initiator",
        )
        if metadata.get(key)
    }

    for control in metadata.get("controls_json", []) or []:
        caption = str(control.get("caption") or "")
        nearby_label = str(control.get("nearby_label") or "")
        value = str(control.get("value") or "")
        checked = bool(control.get("checked"))
        control_metadata = {
            **base_metadata,
            "sheet": control.get("sheet", ""),
            "cell": control.get("cell", ""),
            "caption": caption,
            "checked": checked,
            "value": value,
            "nearby_label": nearby_label,
            "control_source": control.get("source", ""),
        }
        all_chunks.append({
            "file_id": str(doc.id),
            "source_file": doc.original_filename,
            "case_no": case.case_no,
            "page_no": 1,
            "section": str(control.get("sheet") or "controls"),
            "field": caption or str(control.get("cell") or "checkbox"),
            "chunk_index": len(all_chunks),
            "document_type": "staged_excel_control",
            "metadata": control_metadata,
            "text": (
                f"Checkbox control: {nearby_label}\n"
                f"Option: {caption}\n"
                f"Value: {value}\n"
                f"Checked: {'yes' if checked else 'no'}\n"
                f"Cell: {control.get('cell', '')}"
            ).strip(),
        })

    for table in getattr(doc, "tables_json", []) or []:
        headers = [str(header) for header in table.get("headers", [])]
        for row_index, row in enumerate(table.get("rows", []) or []):
            cells = [str(cell) for cell in row]
            row_cells = {
                headers[index] if index < len(headers) else f"column_{index + 1}": cell
                for index, cell in enumerate(cells)
            }
            row_text = " | ".join(cells)
            table_metadata = {
                **base_metadata,
                "table_index": table.get("index", 0),
                "table_caption": table.get("caption", ""),
                "row_index": row_index,
                "headers": headers,
                "cells": row_cells,
            }
            all_chunks.append({
                "file_id": str(doc.id),
                "source_file": doc.original_filename,
                "case_no": case.case_no,
                "page_no": table.get("page_no", 1),
                "section": str(table.get("caption") or "table"),
                "field": "_table_row",
                "chunk_index": len(all_chunks),
                "document_type": "staged_excel_table_row",
                "metadata": table_metadata,
                "text": (
                    f"Table: {table.get('caption', '')}\n"
                    f"Headers: {' | '.join(headers)}\n"
                    f"Row {row_index}: {row_text}"
                ).strip(),
            })

    # Also include full section texts as context chunks
    for sec in doc.sections_json:
        content = sec.get("content", "")
        if content and len(content) > 40:
            all_chunks.append({
                "file_id": str(doc.id),
                "source_file": doc.original_filename,
                "case_no": case.case_no,
                "page_no": sec.get("page_no", 1),
                "section": sec.get("heading", ""),
                "field": "_full_section",
                "chunk_index": len(all_chunks),
                "document_type": f"staged_{doc.file_type}_section",
                "text": content[:2000],
            })

    with open(chunks_path, "wb") as f:
        pickle.dump(all_chunks, f)

    logger.info(
        "Row-level chunks written: %d chunks for %s → %s",
        len(all_chunks), doc.original_filename, chunks_path.name,
    )
    return len(all_chunks)


def _split_text(text: str, chunk_size: int = 800, overlap: int = 120) -> list[str]:
    """Split text into overlapping chunks."""
    if len(text) <= chunk_size:
        return [text] if text.strip() else []

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks
